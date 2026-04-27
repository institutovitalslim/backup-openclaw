#!/usr/bin/env python3
"""
Document Word Creator
Cria documentos .docx formatados com títulos, conteúdo, tabelas e estilos.
"""

import argparse
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document(output_path, title, content, table_data=None, subtitle=None):
    doc = Document()
    
    # Configurar estilos padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    
    # Subtítulo
    if subtitle:
        sub_para = doc.add_paragraph(subtitle)
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.runs[0]
        sub_run.font.size = Pt(12)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()  # Espaço
    
    # Conteúdo
    if content:
        for line in content.split('\n'):
            if line.strip().startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.strip().startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.strip().startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            elif line.strip().startswith('- '):
                doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
            elif line.strip().startswith('1. '):
                doc.add_paragraph(line.replace('1. ', ''), style='List Number')
            else:
                doc.add_paragraph(line)
    
    # Tabela
    if table_data:
        rows = table_data.strip().split('\n')
        if rows:
            cols = rows[0].split(',')
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = 'Table Grid'
            
            # Cabeçalho
            for i, col in enumerate(cols):
                cell = table.rows[0].cells[i]
                cell.text = col.strip()
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            
            # Dados
            for row_data in rows[1:]:
                cells = row_data.split(',')
                row = table.add_row()
                for i, cell_data in enumerate(cells):
                    if i < len(cols):
                        row.cells[i].text = cell_data.strip()
    
    # Rodapé
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("─")
    footer_run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
    
    # Salvar
    doc.save(output_path)
    print(f"✅ Documento criado: {output_path}")
    return output_path

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Criar documento Word')
    parser.add_argument('--output', '-o', required=True, help='Caminho de saída .docx')
    parser.add_argument('--title', '-t', help='Título do documento')
    parser.add_argument('--subtitle', '-s', help='Subtítulo')
    parser.add_argument('--content', '-c', help='Conteúdo (suporta markdown básico)')
    parser.add_argument('--table', help='Dados da tabela (CSV format: cabeçalho\nlinha1\nlinha2)')
    
    args = parser.parse_args()
    
    try:
        create_document(args.output, args.title, args.content, args.table, args.subtitle)
    except Exception as e:
        print(f"❌ Erro: {e}", file=sys.stderr)
        sys.exit(1)
